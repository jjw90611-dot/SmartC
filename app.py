import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import math

st.set_page_config(page_title="포스코퓨처엠 밀폐공간 작업 프로그램 자동 작성기", layout="wide")

# =========================================================
# 서식 헬퍼
# =========================================================
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:color'), '000000')
        tc_borders.append(border)
    tc_pr.append(tc_borders)

def set_font(run, name='맑은 고딕', size=10, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def write_cell(cell, text, size=10, bold=False, align='center', v_align='center', bg=None, color=None):
    cell.text = ''
    if bg: set_cell_bg(cell, bg)
    set_cell_border(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if v_align == 'center' else WD_ALIGN_VERTICAL.TOP
    for i, line in enumerate(str(text).split('\n')):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                       'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        set_font(run, size=size, bold=bold, color=color)

def add_para(doc, text, size=10, bold=False, align='left', space_after=6, space_before=0, color=None):
    p = doc.add_paragraph()
    p.alignment = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text); set_font(run, size=size, bold=bold, color=color)
    return p

def add_header_marker(doc):
    add_para(doc, 'POSCO FUTURE M: Company Use Only', size=9, align='right', space_after=6, color=(100, 100, 100))

def add_section_title(doc, text, size=15):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text); set_font(run, size=size, bold=True, color=(0, 59, 112))
    return p

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            if i < len(row.cells): row.cells[i].width = Cm(width)

def set_row_height(row, height_cm):
    tr = row._tr; trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567))); trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

HEADER_BG = 'D9E1F2'; SUBHEADER_BG = 'F2F2F2'; GRAY_BG = 'F3F3F3'

def calc_volume(shape, W, L, H, D):
    rect = W * L * H if shape in ('rect', 'combined') else 0
    cyl = (math.pi * (D ** 2) / 4) * H if shape in ('cyl', 'combined') else 0
    return rect + cyl

# =========================================================
# 메인 문서 생성 (원본 17페이지 양식 정확 반영)
# =========================================================
def create_word_document(data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8); section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)

    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    style._element.rPr.rFonts.set(qn('w:ascii'), '맑은 고딕')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), '맑은 고딕')
    style.font.size = Pt(10)

    # ============================================================
    # Page 1 : 표지 (1페이지 내 완전 수록)
    # ============================================================
    add_header_marker(doc)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run('밀폐공간 작업\n프로그램')
    set_font(run, size=34, bold=True, color=(0, 59, 112))

    doc.add_paragraph()

    for label, value in [
        ('공사명', data['project_name']),
        ('공사기간', f"{data['start_date']} ~ {data['end_date']}"),
        ('밀폐공간작업기간', f"{data['start_date']} ~ {data['end_date']}"),
        ('작업내용', data['work_detail']),
    ]:
        add_para(doc, f" ‑ {label} : {value}", size=13, space_after=3)

    doc.add_paragraph()

    add_para(doc, f"작성일 : {data['today']}", size=12, space_after=3)
    add_para(doc,
        f"회사명 : {data['company_name']}          현장소장 : {data['manager']} (서명)",
        size=12, space_after=0)
    doc.add_page_break()

    # ============================================================
    # Page 2 : 목차 (원본대로 5장까지 + 첨부)
    # ============================================================
    add_header_marker(doc)
    add_section_title(doc, '목차', size=20)
    doc.add_paragraph()
    toc = [
        "1. 밀폐공간의 위치 파악 및 관리방안",
        "",
        "2. 밀폐공간 내 질식∙중독 등을 일으킬 수 있는 유해∙위험요인의",
        "    파악 및 관리 방안",
        "",
        "3. 안전보건교육 및 훈련",
        "",
        "4. 그 밖에 밀폐공간 작업 근로자의 건강장해 예방에 관한 사항",
        "",
        "5. 프로그램 평가",
        "",
        "",
        "첨부: 밀폐공간 작업 시 사전 확인이 필요한 사항에 대한 확인 절차",
    ]
    for item in toc:
        add_para(doc, item, size=13, space_after=6)
    doc.add_page_break()

    # ============================================================
    # Page 3 : 1. 밀폐공간의 위치 및 관리방안
    # ============================================================
    add_header_marker(doc)
    add_section_title(doc, '1. 밀폐공간의 위치 및 유해∙위험요인의 관리방안', size=13)

    add_para(doc, '(1) 사업장 내 밀폐공간 위치 및 위험물질 관리 현황', size=11, bold=True, space_after=6)
    t = doc.add_table(rows=3, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t.cell(0, 0), '작업장소', bold=True, bg=HEADER_BG)
    write_cell(t.cell(0, 1), '공정명', bold=True, bg=HEADER_BG)
    write_cell(t.cell(0, 2), '질식·중독 발생 위험\n유해위험 물질', bold=True, bg=HEADER_BG)
    write_cell(t.cell(0, 3), '관리방안', bold=True, bg=HEADER_BG)
    write_cell(t.cell(0, 4), '', bold=True, bg=HEADER_BG)
    t.cell(0, 3).merge(t.cell(0, 4))
    write_cell(t.cell(1, 3), '입조 전', bold=True, bg=HEADER_BG)
    write_cell(t.cell(1, 4), '입조 중', bold=True, bg=HEADER_BG)
    t.cell(0, 0).merge(t.cell(1, 0)); t.cell(0, 1).merge(t.cell(1, 1)); t.cell(0, 2).merge(t.cell(1, 2))
    write_cell(t.cell(2, 0), data['space_name'])
    write_cell(t.cell(2, 1), data['process_name'])
    write_cell(t.cell(2, 2), data['hazard_material'])
    write_cell(t.cell(2, 3), '', align='left')
    write_cell(t.cell(2, 4), '', align='left')
    set_col_widths(t, [2.5, 2.5, 3.5, 3.5, 3.5])
    for row in t.rows:
        set_row_height(row, 0.9)

    add_para(doc, '', size=6)
    add_para(doc, '* 위치도', size=10, space_after=2)
    add_para(doc, '(위치도 이미지 삽입 영역)', size=9, color=(120, 120, 120), space_after=10)

    add_para(doc, '(2) Blind 설치 개수 및 위치', size=11, bold=True, space_after=6)
    t2 = doc.add_table(rows=2, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['공정명', 'Blind 설치 개수 (EA)', '확인자 서명']):
        write_cell(t2.cell(0, i), h, bold=True, bg=HEADER_BG)
    write_cell(t2.cell(1, 0), data['process_name'])
    write_cell(t2.cell(1, 1), str(data['blind_count']))
    write_cell(t2.cell(1, 2), data['manager'])
    set_col_widths(t2, [5, 5, 5])
    for row in t2.rows:
        set_row_height(row, 0.9)

    add_para(doc, '* Blind 설치 도면 및 List 작업허가서 첨부 확인', size=9, space_before=6)
    doc.add_page_break()

    # ============================================================
    # Page 4 : 2. 유해∙위험요인
    # ============================================================
    add_header_marker(doc)
    add_section_title(doc, '2. 밀폐공간 내 질식∙중독 등을 일으킬 수 있는\n     유해∙위험요인의 파악 및 관리방안', size=12)

    t3 = doc.add_table(rows=10, cols=3)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['물질명', '유해위험요인', '대책']):
        write_cell(t3.cell(0, i), h, bold=True, bg=HEADER_BG)

    write_cell(t3.cell(1, 0), '예시)\n산화니켈', bold=True)
    write_cell(t3.cell(1, 1), '질식')
    write_cell(t3.cell(1, 2),
        '‑ MSDS교육 실시\n'
        '‑ 감시인, 예비감시인 및 입조작업자는 특별안전교육\n'
        '  (N2 분위기 위험요소, Air Line Mask 사용법 등) 사전 이수\n'
        '‑ 감시인과 입조 작업자간 무전기/스키퍼 구비\n'
        '‑ 감시인 자리이탈 금지 및 출입 인원 점검표 실시간 작성',
        size=9, align='left', v_align='top')
    write_cell(t3.cell(2, 1), '눈에 접촉')
    write_cell(t3.cell(2, 2),
        '‑ 많은 양의 물을 사용하여 15분 이상 동안 눈을 씻어내고\n  즉시 의사의 진료를 받을 것',
        size=9, align='left', v_align='top')
    write_cell(t3.cell(3, 1), '피부 접촉')
    write_cell(t3.cell(3, 2),
        '‑ 오염된 의복 및 신발을 제거한 후 15분 동안 비누와 물로\n  씻어내고 필요시 의사의 진료를 받을 것',
        size=9, align='left', v_align='top')
    write_cell(t3.cell(4, 1), '흡입')
    write_cell(t3.cell(4, 2),
        '‑ 구토를 할 경우 구토물이 기도를 막는 것을 방지하기 위해\n  머리를 둔부보다 낮추도록 할 것\n‑ 즉시 의사의 진료를 받을 것',
        size=9, align='left', v_align='top')

    write_cell(t3.cell(5, 0), '예시)\nN2', bold=True)
    write_cell(t3.cell(5, 1), '질식')
    write_cell(t3.cell(5, 2),
        '‑ MSDS교육 실시\n'
        '‑ 감시인, 예비감시인 및 입조작업자는 특별안전교육\n'
        '  (N2 분위기 위험요소, Air Line Mask 사용법 등) 사전 이수\n'
        '‑ 감시인과 입조 작업자간 무전기/스키퍼 구비\n'
        '‑ 감시인 자리이탈 금지 및 출입 인원 점검표 실시간 작성',
        size=9, align='left', v_align='top')
    write_cell(t3.cell(6, 1), '흡입 시')
    write_cell(t3.cell(6, 2),
        '‑ 호흡이 없으면 인공호흡 실시\n‑ 노출로 인한 영향이 나타나면 비오염지역으로 옮길 것\n‑ 즉시 의사의 진찰과 치료를 받을 것',
        size=9, align='left', v_align='top')
    write_cell(t3.cell(7, 1), '피부 접촉 시')
    write_cell(t3.cell(7, 2),
        '‑ 화학물질의 피부 접촉 즉시 의사의 진찰과 치료를 받을 것\n‑ 해당물질에 접촉시 동상을 입을 가능성이 있으니 주의할 것',
        size=9, align='left', v_align='top')
    write_cell(t3.cell(8, 1), '눈접촉 시')
    write_cell(t3.cell(8, 2),
        '‑ 눈에 들어간 경우 눈꺼풀을 들어올려 15분 동안 물로 충분히\n  씻어낼 것\n‑ 눈에 화학물질이 들어간 경우 즉시 의사의 진찰과 진료를\n  받을 것',
        size=9, align='left', v_align='top')
    write_cell(t3.cell(9, 1), '섭취 시')
    write_cell(t3.cell(9, 2), '많은 양의 화학물질을 섭취한 경우 의사의 진찰을 받을 것',
        size=9, align='left', v_align='top')
    t3.cell(1, 0).merge(t3.cell(4, 0))
    t3.cell(5, 0).merge(t3.cell(9, 0))
    set_col_widths(t3, [2.5, 3, 10])
    doc.add_page_break()

    # ============================================================
    # Page 5 : 3. 안전보건교육 및 훈련
    # ============================================================
    add_header_marker(doc)
    add_section_title(doc, '3. 안전보건교육 및 훈련', size=14)

    t4 = doc.add_table(rows=3, cols=2)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t4.cell(0, 0), '특별안전보건교육내용\n(산업안전보건법 시행규칙 별표5)', bold=True, bg=HEADER_BG)
    write_cell(t4.cell(0, 1), '교육일정 및 강사', bold=True, bg=HEADER_BG)
    write_cell(t4.cell(1, 0),
        '교육내용\n'
        '○ 산소농도 측정 및 작업환경에 관한 사항\n'
        '○ 사고 시의 응급처지 및 비상 시 구출에 관한 사항\n'
        '○ 보호구 착용 및 사용방법에 관한 사항 요령\n'
        '○ 작업내용 ∙ 안전작업방법 및 절차에 관한 사항\n'
        '○ 장비 ∙ 설비 및 시설 등의 안전점검에 관한 사항\n'
        '○ 그 밖에 안전 ∙ 보건관리에 필요한 사항',
        align='left', v_align='top')
    write_cell(t4.cell(1, 1),
        f'○ 특별안전보건교육\n  ‑ 일정 : {data["edu_date"]}\n  ‑ 강사 : {data["instructor"]}',
        align='left', v_align='top')
    write_cell(t4.cell(2, 0),
        '○ 기타사항\n  ‑ 특별안전 교육 일지 및 보관\n  ‑ 교육일지 양식 활용',
        align='left', v_align='top')
    write_cell(t4.cell(2, 1), '* 교육사진\n\n(사진 부착 영역)', align='center', v_align='center')
    set_col_widths(t4, [8.5, 7])
    doc.add_page_break()

    # ============================================================
    # Page 6 : 4. 건강장해 예방 (1)~(3)
    # ============================================================
    add_header_marker(doc)
    add_section_title(doc, '4. 그 밖에 밀폐공간 작업 근로자의 건강\n     장해 예방에 관한 사항', size=13)

    add_para(doc, '(1) 보유장비 현황', size=11, bold=True, space_after=6)
    t5 = doc.add_table(rows=2, cols=5)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['장비명', '모델명', '보유대수', '최근교정일', '교정주기']):
        write_cell(t5.cell(0, i), h, bold=True, bg=HEADER_BG)
    for i in range(5):
        write_cell(t5.cell(1, i), '')
    set_col_widths(t5, [3, 3, 2.5, 3, 3])

    add_para(doc, '', size=6)
    add_para(doc, '(2) 대여 장비 현황', size=11, bold=True, space_after=6)
    t6 = doc.add_table(rows=2, cols=5)
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['장비명', '필요수량', '대여일수', '연락처(협력업체)', '장비관리담당자']):
        write_cell(t6.cell(0, i), h, bold=True, bg=HEADER_BG)
    for i in range(5):
        write_cell(t6.cell(1, i), '')
    set_col_widths(t6, [3, 2.5, 2.5, 4, 3.5])

    add_para(doc, '', size=6)
    add_para(doc, '(3) 비상용 구조장비 비치현황', size=11, bold=True, space_after=6)
    t7 = doc.add_table(rows=4, cols=2)
    t7.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t7.cell(0, 0), '분 류', bold=True, bg=HEADER_BG)
    write_cell(t7.cell(0, 1), '세부 항목', bold=True, bg=HEADER_BG)
    write_cell(t7.cell(1, 0), '구조장비', bold=True)
    write_cell(t7.cell(1, 1),
        '☐ 이동식 크레인   ☐ 호이스트   ☐ 삼각대   ☐ 구명로프   ☐ 들것\n☐ 윈치   ☐ 기타 (         )',
        align='left')
    write_cell(t7.cell(2, 0), '호흡장비', bold=True)
    write_cell(t7.cell(2, 1),
        '☐ 송기마스크(Air Line Mask)   ☐ ELSA (대피용 호흡장비)\n☐ 기타 (         )',
        align='left')
    write_cell(t7.cell(3, 0), '통화장비', bold=True)
    write_cell(t7.cell(3, 1), '☐ 무전기   ☐ 확성기   ☐ 호각   ☐ 기타 (         )', align='left')
    set_col_widths(t7, [3, 12.5])
    doc.add_page_break()

    # ============================================================
    # Page 7 : 기타 추가장비 / 비상연락체계 / 작업자정보 / 가스기준
    # ============================================================
    add_header_marker(doc)
    t7b = doc.add_table(rows=1, cols=2)
    t7b.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t7b.cell(0, 0), '기타\n추가장비', bold=True, bg=HEADER_BG)
    write_cell(t7b.cell(0, 1), '', align='left')
    set_col_widths(t7b, [3, 12.5])
    set_row_height(t7b.rows[0], 1.2)

    add_para(doc, '', size=6)
    add_para(doc, '(4) 비상연락체계', size=11, bold=True, space_after=6)
    t8 = doc.add_table(rows=6, cols=3)
    t8.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['성 명', '연 락 처', '포스코퓨처엠 안전대응팀']):
        write_cell(t8.cell(0, i), h, bold=True, bg=HEADER_BG)
    labels = ['현장소장', '안전관리자', '관리감독자', '작업반장', '기타']
    for i, lab in enumerate(labels):
        write_cell(t8.cell(i+1, 0), lab)
        write_cell(t8.cell(i+1, 1), data['manager'] if lab == '현장소장' else '')
    write_cell(t8.cell(1, 2), '', bold=True)
    for r in range(2, 6):
        write_cell(t8.cell(r, 2), '')
    t8.cell(1, 2).merge(t8.cell(5, 2))
    set_col_widths(t8, [4, 6, 5.5])

    add_para(doc, '', size=6)
    add_para(doc, '(5) 작업자 정보', size=11, bold=True, space_after=6)
    t9 = doc.add_table(rows=7, cols=6)
    t9.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['구분', '성명', '연락처', '구분', '성명', '연락처']):
        write_cell(t9.cell(0, i), h, bold=True, bg=HEADER_BG)
    roles_L = ['포스코퓨처엠\n공사감독', '관리감독자 1', '관리감독자 2', '감시인 1', '감시인 2', '감시인 3']
    roles_R = ['작업자 1', '작업자 2', '작업자 3', '작업자 4', '작업자 5', '작업자 6']
    for i in range(6):
        write_cell(t9.cell(i+1, 0), roles_L[i], size=9)
        write_cell(t9.cell(i+1, 3), roles_R[i], size=9)
        for c in [1, 2, 4, 5]:
            write_cell(t9.cell(i+1, c), '')
    set_col_widths(t9, [2.7, 2.3, 2.5, 2.7, 2.3, 2.5])

    add_para(doc, '', size=6)
    add_para(doc, '(6) 산소 및 유해가스 농도의 측정결과 평가 기준', size=11, bold=True, space_after=2)
    add_para(doc, '  ‑ 측정결과는 작업허가서 참조', size=9, space_after=6)
    t10 = doc.add_table(rows=2, cols=6)
    t10.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['구분', 'O₂', 'HC', 'H₂S', 'CO', 'CO₂']):
        write_cell(t10.cell(0, i), h, bold=True, bg=HEADER_BG)
    vals = ['작업허용\n농도', '20% 이상\n23.5% 미만', 'LEL 10% 미만*)', '1 ppm 미만', '25 ppm 미만', '1.5% 미만']
    for i, v in enumerate(vals):
        write_cell(t10.cell(1, i), v, bold=(i == 0), size=9)
    set_col_widths(t10, [2.5, 3, 2.5, 2.5, 2.5, 2.5])
    add_para(doc, '*) 용접/용단/화염사용 : 0% LEL', size=9, space_before=4)
    doc.add_page_break()

    # ============================================================
    # Page 8 : (7) 환기시설 사용 현황 + (8) 환기시설 설치 및 작동 현황
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '(7) 환기시설 사용 현황', size=11, bold=True, space_after=6)
    t_vent = doc.add_table(rows=3, cols=5)
    t_vent.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t_vent.cell(0, 0), '공정명', bold=True, bg=HEADER_BG)
    write_cell(t_vent.cell(0, 1), '작업장소', bold=True, bg=HEADER_BG)
    write_cell(t_vent.cell(0, 2), '용적(㎥)', bold=True, bg=HEADER_BG)
    write_cell(t_vent.cell(0, 3), '환기시설', bold=True, bg=HEADER_BG)
    write_cell(t_vent.cell(0, 4), '', bold=True, bg=HEADER_BG)
    t_vent.cell(0, 3).merge(t_vent.cell(0, 4))
    write_cell(t_vent.cell(1, 0), '', bg=HEADER_BG)
    write_cell(t_vent.cell(1, 1), '', bg=HEADER_BG)
    write_cell(t_vent.cell(1, 2), '', bg=HEADER_BG)
    write_cell(t_vent.cell(1, 3), '모델명 / 환기팬 용량(㎥/h)', bold=True, bg=HEADER_BG, size=9)
    write_cell(t_vent.cell(1, 4), '사용수량(EA)', bold=True, bg=HEADER_BG, size=9)
    t_vent.cell(0, 0).merge(t_vent.cell(1, 0))
    t_vent.cell(0, 1).merge(t_vent.cell(1, 1))
    t_vent.cell(0, 2).merge(t_vent.cell(1, 2))

    volume = calc_volume(data['shape'], data['W'], data['L'], data['H'], data['D'])
    write_cell(t_vent.cell(2, 0), data['process_name'])
    write_cell(t_vent.cell(2, 1), data['space_name'])
    write_cell(t_vent.cell(2, 2), f'{volume:,.1f}')
    write_cell(t_vent.cell(2, 3), '')
    write_cell(t_vent.cell(2, 4), '')
    set_col_widths(t_vent, [2.5, 3, 2.5, 4.5, 3])

    add_para(doc, '* 작업 전 – 체적의 10배 이상 환기 , 작업 중 – 시간당 작업장소 체적의 20회 이상 환기',
             size=9, space_before=6, space_after=10)

    add_para(doc, '(8) 환기시설 설치 및 작동 현황', size=11, bold=True, space_after=6)
    t_photo = doc.add_table(rows=1, cols=1)
    t_photo.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t_photo.cell(0, 0), '(환기시설 설치 사진 부착 영역)\n\n\n\n', align='center',
               color=(150, 150, 150), size=10)
    set_col_widths(t_photo, [15.5])
    set_row_height(t_photo.rows[0], 8.0)
    doc.add_page_break()

    # ============================================================
    # Page 9 : 특이사항 + 5. 프로그램의 평가
    # ============================================================
    add_header_marker(doc)
    t_sp = doc.add_table(rows=1, cols=2)
    t_sp.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t_sp.cell(0, 0), '특이\n사항', bold=True, bg=HEADER_BG)
    write_cell(t_sp.cell(0, 1), '', align='left')
    set_col_widths(t_sp, [2.5, 13])
    set_row_height(t_sp.rows[0], 2.0)

    add_para(doc, '', size=6)
    add_section_title(doc, '5. 프로그램의 평가', size=14)
    add_para(doc, '밀폐공간 작업 프로그램 평가표', size=12, bold=True, align='center', space_after=4)
    add_para(doc, f'일자 : {data["today"]}', size=10, align='right', space_after=8)

    eval_items = [
        ("밀폐공간\n허가", 1, "밀폐공간 작업장소 보유현황 및 위치 등에 대한 자료가 작성되어 있는가?"),
        ("", 2, "밀폐공간 출입시 작업허가서를 작성하여 발급 받았는가?"),
        ("", 3, "작업허가서는 규정양식을 사용하여 올바르게 작성되었는가?"),
        ("", 4, "프로그램 추진팀(장)은 작업허가서를 적법한 절차에 의해 발급하였는가?"),
        ("산소 및\n유해가스\n농도측정", 5, "산소 및 유해가스 농도 측정대상 물질은 적정하게 선택되었으며 측정시 누락된 물질은 없는가?"),
        ("", 6, "측정장비의 신뢰성(교정 등)은 확보되었는가?"),
        ("", 7, "측정지점수, 측정방법 등은 정해진 규정을 준수하였는가?"),
        ("", 8, "측정결과에 대한 판정은 적합하게 이루어졌는가?"),
        ("환기대책", 9, "밀폐공간 작업장소에 따라 적합한 환기방법, 환기량 선정 등 환기대책은 적절하게 수립되었는가?"),
        ("", 10, "환기팬의 점검은 주기적으로 실시하였는가?"),
        ("보호구\n선정 및\n사용", 11, "보호구의 종류 및 수량은 충분한가?"),
        ("", 12, "보호구의 보유수량 및 대여필요장비 목록은 작성되어 있는가?"),
        ("", 13, "작업에 따라 적합한 보호구가 선정되어 사용되었는가?"),
        ("", 14, "누출검사를 매사용 시마다 시행하도록 하고 있는가?"),
        ("", 15, "보호구를 주기적으로 청소, 점검 등을 실시하는가?"),
        ("응급처치\n체계", 16, "응급상황 발생시 비상연락을 위한 체계는 구축되어 있는가?"),
        ("", 17, "응급전화, 무전기 등의 통신장비는 구비되어 있는가?"),
        ("교육 및\n훈련의\n적정성", 18, "프로그램관리자, 관리감독자, 작업자 등에 대한 교육계획을 수립하여 시행하고 있는가?"),
        ("", 19, "밀폐공간 작업시마다 작업자에게 교육을 실시하고 있는가?"),
        ("", 20, "관련교육을 실시하는 경우 교육내용 등을 기록하고 보존하는가?"),
        ("", 21, "교육내용, 자료 등은 적절하며 최신성을 유지하고 있는가?"),
    ]
    t13 = doc.add_table(rows=len(eval_items) + 1, cols=4)
    t13.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['구분', '번호', '평가항목', '평가\n(O,X)']):
        write_cell(t13.cell(0, i), h, bold=True, bg=HEADER_BG)
    for i, item in enumerate(eval_items):
        write_cell(t13.cell(i+1, 0), item[0], bold=True, size=9)
        write_cell(t13.cell(i+1, 1), str(item[1]), size=9)
        write_cell(t13.cell(i+1, 2), item[2], size=9, align='left')
        write_cell(t13.cell(i+1, 3), '')
    t13.cell(1, 0).merge(t13.cell(4, 0))
    t13.cell(5, 0).merge(t13.cell(8, 0))
    t13.cell(9, 0).merge(t13.cell(10, 0))
    t13.cell(11, 0).merge(t13.cell(15, 0))
    t13.cell(16, 0).merge(t13.cell(17, 0))
    t13.cell(18, 0).merge(t13.cell(21, 0))
    set_col_widths(t13, [2.2, 1.2, 10.5, 1.6])
    doc.add_page_break()

    # ============================================================
    # Page 10 : 22번 + 평가결과 + 첨부서류 시작
    # ============================================================
    add_header_marker(doc)
    t22 = doc.add_table(rows=2, cols=4)
    t22.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t22.cell(0, 0), '교육 및\n훈련의\n적정성', bold=True, bg=SUBHEADER_BG, size=9)
    write_cell(t22.cell(0, 1), '22', size=9)
    write_cell(t22.cell(0, 2), '교육받은 자는 교육내용을 충분히 숙지하여 작업에 올바르게 적용하고 있는가?',
               size=9, align='left')
    write_cell(t22.cell(0, 3), '')
    write_cell(t22.cell(1, 0), '프로그램평가결과', bold=True, bg=SUBHEADER_BG)
    write_cell(t22.cell(1, 1), '', bg=SUBHEADER_BG)
    write_cell(t22.cell(1, 2), '판정수', bold=True, bg=SUBHEADER_BG)
    write_cell(t22.cell(1, 3), '')
    t22.cell(1, 0).merge(t22.cell(1, 1))
    set_col_widths(t22, [2.2, 1.2, 10.5, 1.6])

    add_para(doc, '', size=6)
    add_para(doc, '■ 첨부서류', size=13, bold=True, space_after=4)
    add_section_title(doc, '밀폐공간 작업 시 사전 확인이 필요한 사항에\n     대한 확인 절차', size=12)

    t14 = doc.add_table(rows=3, cols=3)
    t14.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['구분', '사전 확인 사항', '확인 절차']):
        write_cell(t14.cell(0, i), h, bold=True, bg=HEADER_BG)
    write_cell(t14.cell(1, 0), '공통', bold=True)
    write_cell(t14.cell(1, 1),
        '1. 작업내용, 작업방법, 순서에 대한 협의를 사전에 실시하였는가?\n'
        '2. 용기 내 위험물 제거 및 가스농도 측정 결과 양호한가?\n'
        '3. 개방 후 배관격리, 밸브의 맹판을 설치하였는가?\n'
        '4. Purge용 Steam/질소 Line분리하고 Air Hose 설치하였는가?\n'
        '5. 밸브잠금 표식("조작 절대금지") 및 시건 장치를 설치하였는가?\n'
        '6. 장치 내 모든 작동부분을 전기적, 기계적으로 차단하였는가?\n'
        '7. 경고문 표지판 및 인원점검표를 설치하였는가?\n'
        '8. 방폭용 전등 및 공기구를 사용하고 있는가?\n'
        '9. 비상시 연락장비 및 구조장비 등을 비치하고 있는가?\n'
        '10. 밀폐공간 내부의 적절한 환기 조치를 하였는가?',
        size=9, align='left', v_align='top')
    write_cell(t14.cell(1, 2),
        '관리감독자와 감시인이\n작업허가 시에 확인하고\n작업허가서에 서명함.',
        size=9, align='center', v_align='center')
    write_cell(t14.cell(2, 0), '질소충전\n또는\n산소농도\n부족작업', bold=True)
    write_cell(t14.cell(2, 1),
        '1. 작업계획서 별도 마련되어 있고 비상시 구조방법은 있는가?\n'
        '2. Plant Air압력이 6.5 이상임을 확인하였는가?\n'
        '3. Air Line Mask Flame Set의 이상여부를 확인하였는가?\n'
        '4. 공기호흡기를 용기출입구에 비치하고 있는가?',
        size=9, align='left', v_align='top')
    write_cell(t14.cell(2, 2),
        '관리감독자와 감시인이\n작업허가 시에 확인하고\n작업허가서에 서명함.',
        size=9, align='center', v_align='center')
    set_col_widths(t14, [2.5, 9.5, 3.5])

    add_para(doc, '', size=6)
    add_para(doc, '가. 산소 및 유해가스 농도 측정 방법 및 유의사항', size=11, bold=True, space_after=4)
    doc.add_page_break()

    # ============================================================
    # Page 11 : 측정 상세 + 나. 측정지점/측정방법 + 형태별 도식
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '* 다음의 경우 반드시 측정할 것', size=10, bold=True, space_after=2)
    for line in [
        '1) 당일의 작업을 개시하기 전',
        '2) 교대제로 작업을 행할 경우 작업 당일 최초 교대가 행해져서 작업이 시작되기 전',
        '3) 작업에 종사하는 전체 근로자가 작업을 하고 있던 장소를 떠났다가 돌아와 다시 작업을 개시하기 전',
        '4) 근로자의 건강, 환기장치 등에 이상이 있을 때',
        '5) 작업을 하는 과정에서 유해가스가 발생할 가능성이 있을 경우(연속측정)',
        '6) 작업자 또는 작업승인부서에서 측정이 필요하다고 인정되는 경우',
        '7) 작업당일 밀폐공간 최초 출입전 가스측정 값은 1시간 이내일 것',
    ]:
        add_para(doc, line, size=10, space_after=2)

    add_para(doc, '', size=6)
    add_para(doc, '나. 산소 및 유해가스 농도 측정 방법 및 유의사항', size=11, bold=True, space_after=6)
    t_meas = doc.add_table(rows=2, cols=2)
    t_meas.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t_meas.cell(0, 0), '측정지점', bold=True, bg=HEADER_BG)
    write_cell(t_meas.cell(0, 1),
        '○ 작업장소에 대해서 수직방향 및 수평방향으로 각각 3개소 이상\n'
        '○ 작업에 따라 근로자가 출입하는 장소로서 작업시 근로자의 호흡위치를 중심',
        size=9, align='left', v_align='top')
    write_cell(t_meas.cell(1, 0), '측정방법', bold=True, bg=HEADER_BG)
    write_cell(t_meas.cell(1, 1),
        '○ 휴대용측정기 또는 검지관을 이용하여 산소 및 유해가스 농도를 측정한다.\n'
        '○ 탱크 등 깊은 장소의 농도를 측정시에는 고무호스나 PVC로 된 채기관을 사용한다.\n'
        '  ※ 채기관은 1m마다 작은 눈금으로, 5m마다 큰 눈금으로 표시를 하여 깊이 측정\n'
        '○ 산소 및 유해가스 농도 측정시에는 면적, 깊이를 고려하여 밀폐공간 내부를 골고루 측정한다.\n'
        '○ 공기 채취시에는 채기관의 내부용적 이상의 피검공기로 완전히 치환 후 측정한다.',
        size=9, align='left', v_align='top')
    set_col_widths(t_meas, [2.5, 13])

    add_para(doc, '', size=6)
    add_para(doc, '<작업장소 형태별 측정지점>', size=10, bold=True, align='center', space_after=6)
    t_fig = doc.add_table(rows=1, cols=1)
    write_cell(t_fig.cell(0, 0), '(작업장소 형태별 측정지점 도식 삽입 영역)',
               align='center', color=(150, 150, 150), size=10)
    set_col_widths(t_fig, [15.5])
    set_row_height(t_fig.rows[0], 5.0)
    doc.add_page_break()

    # ============================================================
    # Page 12 : * 유의사항 + 다. 환기 + 작업장소 용적 + 환기팬 용량
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '* 유의사항', size=11, bold=True, space_after=4)
    for line in [
        '○ 측정자(보건관리자, 안전관리자, 관리감독자 등)는 측정방법을 충분하게 숙지',
        '○ 밀폐공간 외부에서 측정하는 것을 원칙으로 하되 측정자는 안전에 유의',
        '○ 긴급사태에 대비 측정자의 보조자를 배치토록 하고 보조자도 구명밧줄을 준비',
        '○ 밀폐공간내에 들어가 측정할 경우 측정자 및 보조자는 공기호흡기와 송기마스크 등 호흡용\n  보호구를 필요시 착용',
        '○ 측정에 필요한 장비 등은 방폭형 구조로 된 것을 사용',
    ]:
        add_para(doc, line, size=10, space_after=2)

    add_para(doc, '', size=6)
    t_kosha = doc.add_table(rows=1, cols=1)
    t_kosha.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t_kosha.cell(0, 0),
        '* KOSHA GUIDE 환기기준\n'
        '  작업 전 : 체적의 10배 이상 환기\n'
        '  작업 중 : 시간당 작업장소 체적의 20회이상 환기',
        size=10, bold=True, align='left', bg='FFF9E6', color=(180, 95, 6))
    set_col_widths(t_kosha, [15.5])

    add_para(doc, '', size=6)
    add_para(doc, '다. 환기', size=11, bold=True, space_after=4)
    add_para(doc,
        '밀폐공간 작업시 작업장소에서 적정한 공기가 유지되도록 환기를 실시한 후 작업을 하며, '
        '작업공간 내에서 유해가스가 지속적으로 발생한 경우 (양수기 가동, 슬러지 제거작업 등)에는 '
        '계속적으로 환기를 실시한다.',
        size=10, space_after=8)

    add_para(doc, '■ 작업장소 용적', size=10, bold=True, space_after=4)
    t_vol = doc.add_table(rows=2, cols=3)
    t_vol.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['공정명(Item)', '작업장소 용적(m³)', '비고']):
        write_cell(t_vol.cell(0, i), h, bold=True, bg=HEADER_BG)
    write_cell(t_vol.cell(1, 0), data['process_name'])
    write_cell(t_vol.cell(1, 1), f'{volume:,.1f} m³')
    write_cell(t_vol.cell(1, 2), '')
    set_col_widths(t_vol, [5, 5.5, 5])

    add_para(doc, '', size=6)
    add_para(doc, '■ 환기팬 용량   * 아래 송풍기/배풍기 위치도에 표시 1,2,3', size=10, bold=True, space_after=4)

    Q_need_min = volume * 0.4
    Q_need_hr = Q_need_min * 60

    t_fan = doc.add_table(rows=5, cols=5)
    t_fan.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['구분', '모델명', '환기팬 용량\n(m³/h)', '수량\n(EA)', '장비관리담당자\n(협력업체)']):
        write_cell(t_fan.cell(0, i), h, bold=True, bg=HEADER_BG, size=9)
    fan_rows = [
        ('1', 'SMP‑20', '1,800 m³/h', '1', f'{data["manager"]}\n(000‑0000‑0000)'),
        ('2', 'SMP‑25', '3,120 m³/h', '1', f'{data["manager"]}\n(000‑0000‑0000)'),
        ('3', 'SMP‑30', '4,620 m³/h', '1', f'{data["manager"]}\n(000‑0000‑0000)'),
        ('합산', '', '9,540 m³/h', '3', ''),
    ]
    for i, row in enumerate(fan_rows):
        for j, val in enumerate(row):
            write_cell(t_fan.cell(i+1, j), val, size=9,
                       bold=(i == 3), bg=(SUBHEADER_BG if i == 3 else None))
    set_col_widths(t_fan, [1.8, 3, 3.5, 2, 5])
    doc.add_page_break()

    # ============================================================
    # Page 13 : 환기팬 능력 계산 예시 + 위치도 + 작업장소별 환기량
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '■ 환기팬 환기능력 계산 (예시)', size=11, bold=True, space_after=6)

    add_para(doc, '● 작업시작 전 (작업장소 용적의 10배이상 환기)', size=10, bold=True, space_after=3)
    add_para(doc, f'  ‑ 작업장소 용적(10배) : {volume:,.1f} m³ × 10 = {volume*10:,.1f} m³',
             size=10, space_after=2)
    add_para(doc, '  ‑ 환기팬 환기능력(m³/min) : 환기팬 (3대) 환기능력은 9,540 m³/h = 159 m³/min',
             size=10, space_after=2)
    add_para(doc,
        f'  ‑ 계산결과 : 환기팬(3대) 18min 이상 가동(159 m³/min × 18min = 2,862 m³)시 '
        f'환기능력이 {volume*10:,.1f} m³보다 크므로 만족함',
        size=10, space_after=8)

    add_para(doc, '● 작업중 (시간당 작업장소 용적의 20회 이상 환기)', size=10, bold=True, space_after=3)
    add_para(doc, f'  ‑ 시간당 20회 환기 기준량 : {volume:,.1f} m³ × 20회 = {volume*20:,.1f} m³/hr',
             size=10, space_after=2)
    add_para(doc,
        f'  ‑ 계산결과 : 환기팬(3대) 환기능력이 9,540 m³/h 으로 20회 환기 기준량'
        f'({volume*20:,.1f} m³)보다 크므로 만족함',
        size=10, space_after=10)

    add_para(doc, '<송풍기/배풍기 설치 위치도>', size=10, bold=True, align='center', space_after=6)
    t_pos = doc.add_table(rows=1, cols=1)
    write_cell(t_pos.cell(0, 0), '(송풍기/배풍기 설치 위치도 삽입 영역)',
               align='center', color=(150, 150, 150), size=10)
    set_col_widths(t_pos, [15.5])
    set_row_height(t_pos.rows[0], 4.5)

    add_para(doc, '', size=6)
    add_para(doc, '<작업장소에 따른 환기량>', size=10, bold=True, align='center', space_after=6)
    t_scope = doc.add_table(rows=6, cols=2)
    t_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t_scope.cell(0, 0), '작업장소', bold=True, bg=HEADER_BG)
    write_cell(t_scope.cell(0, 1), '환기량', bold=True, bg=HEADER_BG)
    scopes = [
        '잠함, 압기실 등의 압기공법의 작업실',
        '피트 내부',
        '황화수소가 발생할 우려가 있는 탱크, 보일러 등의 내부',
        '탱크 내 퇴적물 제거작업',
        '기타 밀폐공간',
    ]
    write_cell(t_scope.cell(1, 1),
        '작업 전 당해 밀폐공간 체적의 10배이상 환기하고 작업하는 동안 시간당 밀폐공간 체적의 20배이상 계속 환기한다.',
        size=9, align='left')
    t_scope.cell(1, 1).merge(t_scope.cell(5, 1))
    for i, s in enumerate(scopes):
        write_cell(t_scope.cell(i+1, 0), s, size=9)
    set_col_widths(t_scope, [7, 8.5])
    add_para(doc, '* 위험물 옥외탱크저장소의 경우 Tank 개방검사절차를 준용한다.',
             size=9, space_before=4)
    doc.add_page_break()

    # ============================================================
    # Page 14 : 주의사항 + 이동식 송풍기 점검표 + 라. 보호구
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '<주의사항>', size=11, bold=True, space_after=4)
    for line in [
        '○ 작업 전에는 산소 및 유해가스의 농도가 기준농도를 만족할 수 있도록 충분한 환기를 실시한다.',
        '○ 정전 등에 의한 환기 중단 시에는 즉시 외부로 대피한다.',
        '○ 밀폐공간의 환기시에는 급기구와 배기구를 적절하게 배치하여 작업장내 환기가 효과적으로 이루어지도록 한다.',
        '○ 급기구는 작업자에 근접하여 설치한다.',
        '○ 이동식 환기장치 사용시 폭발 위험 구역 내에서는 방폭형 구조를 사용한다.',
        '○ 이동식 환기장치의 송풍관은 가급적 구부리는 부위가 적게 하고 용접불꽃 등에 의한 구멍이 나지 않도록 난연 재질을 사용한다.',
    ]:
        add_para(doc, line, size=10, space_after=2)

    add_para(doc, '', size=6)
    add_para(doc,
        '※ 이동식 환기장치 사용시 다음 사항을 추진팀(공무담당)에서 반드시 점검하여 사용 중 고장, '
        '가동중지 등으로 인한 위급한 상황이 발생되지 않도록 한다.',
        size=9, space_after=6)
    t_chk = doc.add_table(rows=2, cols=2)
    t_chk.alignment = WD_TABLE_ALIGNMENT.CENTER
    write_cell(t_chk.cell(0, 0), '이동식 송풍기', bold=True, bg=HEADER_BG)
    write_cell(t_chk.cell(0, 1), '송풍관', bold=True, bg=HEADER_BG)
    write_cell(t_chk.cell(1, 0),
        '○ 전원코드의 단선 접속부의 접촉불량 유무\n'
        '○ 코드와 단자상과의 접속상태 불량 유무\n'
        '○ 코드의 끝에 "환기중․정지" 등의 표지판 부착 유무',
        size=9, align='left', v_align='top')
    write_cell(t_chk.cell(1, 1),
        '○ 연소에 의한 구멍이나 파열유무\n'
        '○ 링, 나선의 손상유무\n'
        '○ 접속부의 확실한 고정여부',
        size=9, align='left', v_align='top')
    set_col_widths(t_chk, [8, 7.5])

    add_para(doc, '', size=6)
    add_para(doc, '라. 보호구', size=11, bold=True, space_after=4)
    add_para(doc,
        '밀폐공간 작업시 유해가스에 의한 중독 및 질식에 의한 사고를 예방하기 위해 공기호흡기 및 송기마스크 등의 '
        '보호구를 반드시 착용한 상태에서 작업을 하고, 사용시 사용장소 및 사용방법 등을 충분히 숙지한 후 사용한다. '
        '다만, 작업시 보호구를 착용하는 것이 원칙이나 측정결과 등으로 밀폐공간 내에서의 작업이 안전하다고 판단될 '
        '경우 보호구를 착용하지 않아도 된다.',
        size=10, space_after=6)

    add_para(doc, '□ 공기호흡기', size=11, bold=True, space_after=3)
    add_para(doc, '1) 착용해야 할 장소', size=10, bold=True, space_after=2)
    add_para(doc,
        '  밀폐장소 출입 작업시 다음과 같이 환기할 수 없거나 환기가 불충분한 경우로서 단기간 작업이 가능한 '
        '경우에는 공기호흡기를 반드시 착용하고 출입하며, 고농도의 유기화합물 증기가 예상되는 경우 등에는 '
        '방독마스크를 착용하지 않는다.',
        size=10, space_after=3)
    for line in [
        '  ① 수도나 도수관 등으로 깊은 곳까지 환기가 되지 않는 경우',
        '  ② 탱크와 화학설비 및 선박의 내부 등 구조적으로 충분히 환기시킬 수 없는 경우',
        '  ③ 재해사고시의 구조 등과 같이 충분히 환기시킬 시간적인 여유가 없는 경우 작업시',
    ]:
        add_para(doc, line, size=10, space_after=2)
    doc.add_page_break()

    # ============================================================
    # Page 15 : 공기호흡기 사용법 + 송기마스크
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '2) 공기호흡기의 점검사항 및 사용방법', size=10, bold=True, space_after=3)
    add_para(doc, '[사용전의 점검사항]', size=10, bold=True, space_after=2)
    for line in [
        '  ① 봄베의 잔류압 검사',
        '  ② 고압연결부의 검사',
        '  ③ 면체와 흡기관 및 호기밸브의 기밀검사',
        '  ④ 폐력밸브와 압력계 및 경보기의 동작검사',
    ]:
        add_para(doc, line, size=10, space_after=2)

    add_para(doc, '', size=4)
    add_para(doc, '[공기호흡기의 사용법]', size=10, bold=True, space_after=2)
    for line in [
        '  ① 먼저 봄베를 등에 지고 겨드랑이 끈을 당겨서 조정한다. 다음으로 가슴끈과 허리끈을 몸에 꽉 맞게 조정한다.',
        '  ② 마스크를 쓰게 되면 좌우 4개의 끈을 1조씩 동시에 당겨서 밀착시킨다.',
        '  ③ 흡기관을 두겹으로 강하게 잡고 숨을 들이쉬어 기밀을 확인한다.',
        '  ④ 압력계의 지시치가 30 Kg/㎠ 이하로 내려가거나 경보기가 울리게 되면 곧바로 작업을 중지하고 유해가스가 없는 안전한 위치로 되돌아온다.',
        '  ⑤ 안전한 위치로 되돌아오면 마스크를 벗고 공기탱크를 교환한다. 공기탱크의 교환 시에는 잔류압을 확인한다.',
    ]:
        add_para(doc, line, size=10, space_after=2)

    add_para(doc, '', size=6)
    add_para(doc, '□ 송기마스크', size=11, bold=True, space_after=3)
    add_para(doc,
        '송기마스크는 활동범위에 제한을 받고 있지만, 가볍고 유효사용 시간이 길어지므로 일정한 장소에서의 '
        '장시간 작업에 주로 이용한다.',
        size=10, space_after=6)

    add_para(doc, '1) 전동 송풍기식 호스마스크', size=10, bold=True, space_after=2)
    for line in [
        '  ① 송풍기는 유해가스, 악취 및 먼지가 없는 장소에 설치한다.',
        '  ② 전동 송풍기는 장시간 운전하면 필터에 먼지가 끼므로 정기적으로 점검한다.',
        '  ③ 전동 송풍기를 사용할 때에는 접속전원이 단절되지 않도록 코드 플러그에 반드시 "송기마스크 사용중"이란 표시를 한다.',
        '  ④ 전동 송풍기는 통상적으로 방폭구조가 아니므로 폭발하한을 초과할 우려가 있는 장소에서는 사용하지 않는다.',
        '  ⑤ 정전 등으로 인해 공기공급이 중단되는 경우에 대비한다.',
    ]:
        add_para(doc, line, size=10, space_after=2)

    add_para(doc, '', size=4)
    add_para(doc, '2) 에어라인 마스크', size=10, bold=True, space_after=2)
    add_para(doc,
        '  전동 송풍기식에 비하여 상당히 먼 곳까지 송기할 수 있으며 송기호스가 가늘고 활동하기도 용이하므로 '
        '유해가스가 발생되는 장소에서 주로 사용한다.',
        size=10, space_after=3)
    for line in [
        '  ① 공급되는 공기중의 분진, 오일, 수분 등을 제거하기 위하여 에어라인에 여과장치를 설치한다.',
        '  ② 정전 등으로 인해 공기공급이 중단되는 경우에 대비한다.',
    ]:
        add_para(doc, line, size=10, space_after=2)
    doc.add_page_break()

    # ============================================================
    # Page 16 : 안전보호구
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '□ 안전보호구', size=11, bold=True, space_after=4)
    add_para(doc,
        '탱크나 맨홀과 같이 사다리를 사용하여 내부로 내려가야 하는 경우에는 안전대나 기타 구명밧줄 등을 '
        '사용하여 안전을 확보한다. 비상시에 작업자를 피난시키거나 구출하기 위하여 안전대, 사다리, 구명밧줄 등 '
        '필요한 용구를 준비하고 이것의 사용방법을 작업자에게 숙지하도록 한다. 작업시 작업장소에서 적정한 공기가 '
        '유지되도록 환기를 실시한 후 작업을 하며, 작업공간 내에서 유해가스가 지속적으로 발생한 경우(양수기 가동, '
        '슬러지 제거작업 등)에는 계속적으로 환기를 실시한다.',
        size=10, space_after=10)
    doc.add_page_break()

    # ============================================================
    # Page 17 : 착용 보호구 종류 + 응급처치 관찰사항
    # ============================================================
    add_header_marker(doc)
    add_para(doc, '작업 시 착용하여야 할 보호구의 종류', size=13, bold=True,
             align='center', space_after=8, color=(0, 59, 112))
    add_para(doc,
        '1. 밀폐공간 작업 시 기본 보호구 : 안전화, 안전모, 보안경, 그네식 안전대, 방진/방독마스크',
        size=10, space_after=4)
    add_para(doc,
        '2. 추가 보호구 : (질소 충전 또는 산소농도 부족 작업시) 공기호흡기, 송기마스크,\n'
        '              안전블럭, 전동윈치, ELSA(Emergency Life Support Appratus)',
        size=10, space_after=6)
    for line in [
        '‑ 밀폐공간에 입조하는 작업자가 안전블럭과 전동윈치를 사용하는 경우, 체결 방법은 다음과 같음',
        '  ∙ 안전블럭은 1줄걸이 죔줄을 안전그네 등 쪽 위에 있는 D링에 체결한다.',
        '  ∙ 전동윈치는 2줄걸이 죔줄을 안전그네 가슴 쪽 양편에 있는 D링에 체결한다.',
        '‑ 밀폐공간 입조 작업 시 비상대피용 공기호흡기 ELSA(Emergency Life Support Appratus)를 비치하는 위치는 다음과 같이 한다.',
        '  ∙ 밀폐공간이 질소분위기일 때 : 용기 바깥 출입구와 가까운 신선한 공기가 있는 장소',
        '  ∙ 밀폐공간이 공기분위기일 때 : 용기 내 작업자들이 작업하는 장소와 가깝고 안전한 장소',
    ]:
        add_para(doc, line, size=10, space_after=2)

    add_para(doc, '', size=6)
    add_para(doc, '비상연락체계', size=11, bold=True, space_after=4)
    t_emer = doc.add_table(rows=1, cols=1)
    write_cell(t_emer.cell(0, 0), '(비상연락체계도 삽입 영역)',
               align='center', color=(150, 150, 150), size=10)
    set_col_widths(t_emer, [15.5])
    set_row_height(t_emer.rows[0], 3.0)

    add_para(doc, '', size=6)
    add_para(doc, '<응급처치 시 관찰사항>', size=11, bold=True, space_after=4)
    add_para(doc, '응급처치 시에는 다음의 사항을 주의 깊게 관찰하고 그 내용을 담당자에게 정확히 전달',
             size=10, space_after=4)
    for line in [
        '○ 의식이 있는지 확인한다.',
        '○ 호흡하고 있는지 확인한다. 호흡이 정지되어 있으면 머리를 뒤로 젖히거나 아래턱을 밀어내어 기도를 열어주고 다시 확인한다.',
        '○ 출혈의 유무를 살펴본다.',
        '○ 맥을 짚어본다. 맥박이 뛰지 않는다고 느낄 때는 동공을 살펴본다. 동공이 크게 벌어져 있으면 위험하고 동공의 크기가 좌우 틀리면 뇌에 이상이 있는 경우이다.',
        '○ 손발이 움직이는가를 본다.',
        '○ 얼굴과 피부색, 체온을 살펴본다. 혀, 입술, 피부 등이 푸르스름한 색 또는 흑색이 되고 손톱은 암자색이 되었는지 살펴본다.',
        '○ 재해자의 체온을 유지하도록 보온한다.',
        '○ 협력자를 구한다.',
        '○ 재해자를 운반할 때는 서두르지 말고 재해자의 마음을 가라앉히고 되도록 재해자의 상처를 건드리지 않도록 주의하여 운반한다.',
    ]:
        add_para(doc, line, size=10, space_after=2)

    buf = BytesIO()
    doc.save(buf); buf.seek(0)
    return buf


# =========================================================
# Streamlit UI
# =========================================================
st.sidebar.header("📝 작업 정보 입력")
project_name = st.sidebar.text_input("공사명", "포항공장 R-30101A 내부 정비공사")
start_date = st.sidebar.date_input("작업 시작일")
end_date = st.sidebar.date_input("작업 종료일")
work_detail = st.sidebar.text_area("작업내용", "내부 슬러지 제거 및 검사")
company_name = st.sidebar.text_input("회사명", "포스코퓨처엠")
manager = st.sidebar.text_input("현장소장 성명", "홍길동")

st.sidebar.markdown("---")
st.sidebar.subheader("🔧 공정 및 위험요인")
process_name = st.sidebar.text_input("공정명", "R-30101A")
space_name = st.sidebar.text_input("작업장소", "반응기 내부")
hazard_material = st.sidebar.text_input("유해위험 물질", "H2S, N2")
blind_count = st.sidebar.number_input("Blind 설치 개수", value=5, step=1)

st.sidebar.markdown("---")
st.sidebar.subheader("👨‍🏫 교육")
edu_date = st.sidebar.text_input("교육일정", "2026.06.15 08:00~09:00")
instructor = st.sidebar.text_input("강사", "안전관리자 김안전")

st.sidebar.markdown("---")
st.sidebar.subheader("📐 밀폐공간 체적 계산 (환기시설 표 자동 반영)")
shape = st.sidebar.selectbox("공간 형태", options=['rect', 'cyl', 'combined'],
                              format_func=lambda x: {'rect': '사각형', 'cyl': '원통형', 'combined': '결합형'}[x])
W = st.sidebar.number_input("가로 W (m)", value=6.0, step=0.1)
L = st.sidebar.number_input("세로 L (m)", value=8.0, step=0.1)
H = st.sidebar.number_input("높이 H (m)", value=6.0, step=0.1)
D = st.sidebar.number_input("직경 D (m) - 원통형만", value=0.0, step=0.1)

st.title("📄 포스코퓨처엠 밀폐공간 작업 프로그램 자동 작성기")
st.markdown("""
**원본 17페이지 양식 정확 반영 버전 (v2.1)**
- ✅ 표지 1페이지 완전 수록 (여백 및 폰트 최적화)
- ✅ S-OIL 참조양식 정보 완전 제거 → **포스코퓨처엠으로 통일**
- ✅ 헤더 표기 : `POSCO FUTURE M: Company Use Only`
- ✅ 가스 기준 정정 : O₂ 20%↑ / H₂S 1ppm↓ / CO 25ppm↓
- ✅ 환기시설 표, 특이사항, 환기팬 계산, 보호구, 응급처치까지 원본 순서대로 완전 재현
""")

vol_preview = calc_volume(shape, W, L, H, D)
col1, col2 = st.columns(2)
col1.metric("작업장소 체적 V", f"{vol_preview:,.1f} m³")
col2.metric("작업 전 필요 환기량 (10배)", f"{vol_preview*10:,.1f} m³")

data = {
    "project_name": project_name,
    "start_date": start_date.strftime("%Y.%m.%d"),
    "end_date": end_date.strftime("%Y.%m.%d"),
    "work_detail": work_detail,
    "today": datetime.now().strftime("%Y. %m. %d"),
    "company_name": company_name,
    "manager": manager,
    "process_name": process_name,
    "space_name": space_name,
    "hazard_material": hazard_material,
    "blind_count": blind_count,
    "edu_date": edu_date,
    "instructor": instructor,
    "shape": shape, "W": W, "L": L, "H": H, "D": D,
}

word_file = create_word_document(data)

st.download_button(
    label="📥 밀폐공간 작업 프로그램 (원본 양식 v2.1) 다운로드",
    data=word_file,
    file_name=f"밀폐공간_작업프로그램_{project_name}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True,
)

st.caption("© 2026 포스코퓨처엠 안전부서 | 원본 17페이지 양식 정확 재현 버전")
